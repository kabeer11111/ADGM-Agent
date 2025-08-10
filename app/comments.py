from docx import Document
from docx.enum.text import WD_COLOR_INDEX


def add_safe_heading(doc, text, level=1):
    try:
        doc.add_heading(text, level=level)
    except KeyError:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        

def insert_comments(doc_path, issues, summary_text=None):
    doc = Document(doc_path)

    # === Add Summary at top ===
    if summary_text:
        doc.paragraphs[0].insert_paragraph_before("=== Document Summary ===", style="Heading 1")
        doc.paragraphs[0].insert_paragraph_before(summary_text, style="Normal")
        doc.paragraphs[0].insert_paragraph_before("")  # Blank line

    # === Add inline highlighted comments ===
    for issue in issues:
        try:
            para_num = int(issue["section"].split()[-1]) - 1
        except:
            para_num = -1
        para = doc.paragraphs[para_num] if para_num >= 0 else doc.paragraphs[-1]

        suggestion = issue.get("suggestion", "").strip()
        citations = ", ".join(issue.get("citations", []))
        comment_text = f"[Comment: {suggestion}"
        if citations:
            comment_text += f" | Law: {citations}"
        comment_text += "]"

        
        run = para.add_run(" " + comment_text)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # highlight comment text

    reviewed_path = doc_path.replace(".docx", "_reviewed.docx")
    doc.save(reviewed_path)
    return reviewed_path
