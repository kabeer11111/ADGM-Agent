from docx import Document
from app.models.llm_client import GeminiClient

gemini = GeminiClient()

def extract_meaningful_content(doc_path):
    """Extract the most meaningful content from document for summarization"""
    try:
        doc = Document(doc_path)
        
        # Get all paragraph text
        paragraphs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text and len(text) > 10:  # Skip very short paragraphs
                paragraphs.append(text)
        
        # Get table content
        table_content = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_content.append(" | ".join(row_text))
        
        # Combine content
        all_content = paragraphs + table_content
        full_text = "\n".join(all_content)
        
        # If content is too long, prioritize the beginning and end
        if len(full_text) > 4000:
            # Take first 2000 characters and last 1000 characters
            beginning = full_text[:2000]
            ending = full_text[-1000:]
            full_text = beginning + "\n...\n" + ending
        
        return full_text, len(paragraphs), len(table_content)
        
    except Exception as e:
        print(f"Error extracting content from {doc_path}: {e}")
        return "", 0, 0

def summarize_document(doc_path, doc_type=None):
    """
    Generate a comprehensive summary of any ADGM document
    """
    try:
        # Extract document content
        content, paragraph_count, table_count = extract_meaningful_content(doc_path)
        
        if not content.strip():
            return "[Document appears to be empty or unreadable]"
        
        # Create doc-type specific prompt
        if doc_type:
            type_specific_prompt = f"This is a {doc_type}. "
        else:
            type_specific_prompt = "This appears to be an ADGM legal document. "
        
        prompt = f"""
        {type_specific_prompt}Provide a concise summary with the following structure:

        **Purpose**: What is this document for?
        **Key Elements**: List 3-5 main components/clauses
        **Parties Involved**: Who are the parties (if applicable)?
        **Important Dates**: Any key dates mentioned
        **Compliance Notes**: Any ADGM-specific requirements or references

        Keep each section to 1-2 sentences maximum. Focus on practical business information.
        
        Document content:
        """
        
        # Get summary from LLM
        summary = gemini.ask(prompt, content)
        
        if not summary or summary.strip() in ["", "[No suggestion available]", "[No response"]:
            # Fallback summary
            return f"""
            **Purpose**: {doc_type or 'ADGM Document'} ({paragraph_count} sections, {table_count} tables)
            **Key Elements**: Document contains standard legal provisions and clauses
            **Status**: Requires manual review for detailed analysis
            **Note**: Automated summary unavailable - please review document manually
            """
        
        # Add document stats
        stats = f"\n\n*Document Statistics: {paragraph_count} paragraphs, {table_count} tables*"
        
        return summary.strip() + stats
        
    except Exception as e:
        print(f"Error summarizing document {doc_path}: {e}")
        return f"""
        **Error**: Unable to generate summary ({str(e)[:50]}...)
        **Recommendation**: Please check document format and content manually
        **File**: {doc_path.split('/')[-1] if '/' in doc_path else doc_path}
        """

def summarize_multiple_documents(file_paths_and_types):
    """
    Generate summaries for multiple documents
    Returns dict of {filepath: summary}
    """
    summaries = {}
    
    for file_path, doc_type in file_paths_and_types:
        try:
            summary = summarize_document(file_path, doc_type)
            summaries[file_path] = summary
            print(f"Generated summary for {file_path} ({doc_type})")
        except Exception as e:
            summaries[file_path] = f"[Summary error: {str(e)[:100]}]"
            print(f"Failed to summarize {file_path}: {e}")
    
    return summaries

def get_document_overview(summaries_dict, detected_process=None):
    """
    Create an overview of all documents
    """
    if not summaries_dict:
        return "No documents processed."
    
    overview = f"## Document Package Overview\n"
    if detected_process:
        overview += f"**Process Type**: {detected_process}\n"
    
    overview += f"**Total Documents**: {len(summaries_dict)}\n\n"
    
    for i, (file_path, summary) in enumerate(summaries_dict.items(), 1):
        filename = file_path.split('/')[-1] if '/' in file_path else file_path
        overview += f"### {i}. {filename}\n{summary}\n\n"
    
    return overview
