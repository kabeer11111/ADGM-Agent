import regex as re
from docx import Document
from app.models.llm_client import GeminiClient

# Improved TYPE_RULES with more specific patterns and keywords
TYPE_RULES = [
    ("Articles of Association", [
        r"\bArticles\s+of\s+Association\b",
        r"\bA\s*O\s*A\b",
        r"Article\s+\d+",
        r"company\s+constitution",
        r"articles\s+and\s+regulations",
        r"corporate\s+articles"
    ]),
    ("Memorandum of Association", [
        r"\bMemorandum\s+of\s+Association\b",
        r"\bM\s*O\s*A\b",
        r"objects\s+of\s+the\s+company",
        r"company\s+objects",
        r"memorandum\s+and\s+articles"
    ]),
    ("Board Resolution", [
        r"\bBoard\s+Resolution\b",
        r"\bResolution\s+of\s+(?:the\s+)?Board\b",
        r"Directors?\s+Resolution",
        r"Board\s+of\s+Directors\s+Resolution",
        r"resolved\s+that",
        r"IT\s+IS\s+RESOLVED",
        r"board\s+meeting",
        r"directors?\s+meeting"
    ]),
    ("Shareholder Resolution", [
        r"\bShareholder\s+Resolution\b",
        r"\bMember\s+Resolution\b",
        r"Resolution\s+of\s+(?:the\s+)?Shareholders?",
        r"Resolution\s+of\s+(?:the\s+)?Members?",
        r"shareholders?\s+meeting",
        r"members?\s+meeting",
        r"general\s+meeting"
    ]),
    ("Employment Contract", [
        r"\bEmployment\s+Contract\b",
        r"\bContract\s+of\s+Employment\b",
        r"\bEmployment\s+Agreement\b",
        r"terms\s+and\s+conditions\s+of\s+employment",
        r"employee\s+handbook",
        r"job\s+description",
        r"salary\s+and\s+benefits",
        r"working\s+hours",
        r"probation\s+period",
        r"termination\s+clause"
    ]),
    ("UBO Declaration", [
        r"\bUBO\s+Declaration\b",
        r"\bUltimate\s+Beneficial\s+Owner\b",
        r"beneficial\s+ownership",
        r"UBO\s+information",
        r"ownership\s+structure",
        r"controlling\s+interest"
    ]),
    ("Register of Members and Directors", [
        r"\bRegister\s+of\s+Members\b",
        r"\bRegister\s+of\s+Directors\b",
        r"Members?\s+Register",
        r"Directors?\s+Register",
        r"shareholding\s+details",
        r"director\s+details",
        r"member\s+information"
    ]),
    ("Incorporation Application Form", [
        r"\bIncorporation\s+Application\b",
        r"Application\s+for\s+Incorporation",
        r"company\s+registration",
        r"incorporation\s+form",
        r"registration\s+application"
    ]),
    ("Change of Registered Address Notice", [
        r"Change\s+of\s+(?:Registered\s+)?Address",
        r"Address\s+Change\s+Notice",
        r"registered\s+office\s+change",
        r"change\s+of\s+registered\s+office"
    ]),
    ("Licensing Regulatory Filing", [
        r"Licensing\s+(?:Regulatory\s+)?Filing",
        r"License\s+Application",
        r"regulatory\s+filing",
        r"license\s+renewal",
        r"permit\s+application"
    ]),
    ("Commercial Agreement", [
        r"\bCommercial\s+Agreement\b",
        r"\bService\s+Agreement\b",
        r"business\s+agreement",
        r"commercial\s+contract",
        r"service\s+contract",
        r"supply\s+agreement"
    ]),
    ("Compliance Risk Policy", [
        r"\bCompliance\s+Policy\b",
        r"\bRisk\s+Policy\b",
        r"compliance\s+procedures",
        r"risk\s+management",
        r"internal\s+controls"
    ]),
    ("Renewal Application Form", [
        r"Renewal\s+Application",
        r"License\s+Renewal",
        r"permit\s+renewal",
        r"registration\s+renewal"
    ]),
    ("Compliance Declaration", [
        r"Compliance\s+Declaration",
        r"regulatory\s+compliance",
        r"compliance\s+statement",
        r"declaration\s+of\s+compliance"
    ])
]

def extract_document_content(filepath):
    """Extract comprehensive text content from document"""
    try:
        doc = Document(filepath)
        
        # Get paragraph text
        paragraph_text = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraph_text.append(p.text.strip())
        
        # Get table text
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(" | ".join(row_text))
        
        # Combine all text
        all_text = paragraph_text + table_text
        combined_text = "\n".join(all_text)
        
        return combined_text, paragraph_text, table_text
    
    except Exception as e:
        print(f"Error extracting content from {filepath}: {e}")
        return "", [], []

def detect_doc_type(filepath):
    """Improved document type detection with scoring system"""
    
    combined_text, paragraph_text, table_text = extract_document_content(filepath)
    
    if not combined_text.strip():
        return "Unknown"
    
    # Score each document type based on keyword matches
    type_scores = {}
    
    for doc_type, patterns in TYPE_RULES:
        score = 0
        matches = []
        
        for pattern in patterns:
            # Count matches in the text
            if isinstance(pattern, str) and not pattern.startswith('\\b'):
                # Simple string search
                pattern_matches = len(re.findall(pattern, combined_text, flags=re.I))
            else:
                # Regex pattern
                pattern_matches = len(re.findall(pattern, combined_text, flags=re.I))
            
            if pattern_matches > 0:
                score += pattern_matches
                matches.append(pattern)
        
        if score > 0:
            type_scores[doc_type] = {
                'score': score,
                'matches': matches
            }
    
    # If we have clear matches, return the highest scoring type
    if type_scores:
        best_match = max(type_scores.items(), key=lambda x: x[1]['score'])
        print(f"Document type detection for {filepath}:")
        print(f"Best match: {best_match[0]} (score: {best_match[1]['score']})")
        print(f"Matching patterns: {best_match[1]['matches']}")
        return best_match[0]
    
    # Fallback to LLM if no clear pattern matches
    print(f"No pattern matches found for {filepath}, using LLM fallback")
    try:
        gemini = GeminiClient()
        
        # Create a more detailed prompt with examples
        doc_types_list = [doc_type for doc_type, _ in TYPE_RULES]
        
        prompt = f"""
        Analyze the following document content and classify it as ONE of these document types:
        {', '.join(doc_types_list)}
        
        Look for key indicators like:
        - Document title and headers
        - Legal language and clauses
        - Specific terminology
        - Document structure and format
        
        Respond with ONLY the document type name, nothing else.
        
        Document content (first 2000 characters):
        """
        
        response = gemini.ask(prompt, combined_text[:2000])
        
        # Clean and validate the response
        if response and response.strip():
            response = response.strip().replace('"', '').replace("'", "")
            
            # Check if the response matches one of our known types
            for doc_type, _ in TYPE_RULES:
                if doc_type.lower() in response.lower():
                    print(f"LLM classified as: {doc_type}")
                    return doc_type
        
        print(f"LLM response '{response}' didn't match known types")
        
    except Exception as e:
        print(f"Error using LLM fallback: {e}")
    
    return "Unknown"

def get_document_keywords(doc_type):
    """Get keywords associated with a document type for debugging"""
    for dtype, patterns in TYPE_RULES:
        if dtype == doc_type:
            return patterns
    return []