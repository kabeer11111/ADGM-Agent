import os
import pdfplumber
from docx import Document
import chromadb
from sentence_transformers import SentenceTransformer

PDF_DIR = "data/adgm_laws/"
DB_PATH = "./chroma_db"

# Use persistent client - CRITICAL FIX
client = chromadb.PersistentClient(path=DB_PATH)
db = client.get_or_create_collection("adgm_laws")
embed_model = SentenceTransformer("all-MiniLM-L6-v2")

def extract_pdf_text(pdf_path):
    """Extract all text from PDF file."""
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += (page.extract_text() or "") + "\n"
    return text

def extract_docx_text(docx_path):
    """Extract all text from DOCX file."""
    doc = Document(docx_path)
    full_text = []
    for p in doc.paragraphs:
        if p.text.strip():
            full_text.append(p.text)
    # Include table text too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    return "\n".join(full_text)

def ingest():
    # Check if collection already has data
    global db
    try:
        existing_count = db.count()
        if existing_count > 0:
            print(f"Database already contains {existing_count} documents.")
            response = input("Do you want to clear and re-ingest? (y/N): ")
            if response.lower() == 'y':
                client.delete_collection("adgm_laws")

                db = client.get_or_create_collection("adgm_laws")
                print("Cleared existing database.")
            else:
                print("Keeping existing data.")
                return
    except:
        pass  # Collection might not exist yet
    
    # Support both PDF and DOCX
    files = [
        os.path.join(PDF_DIR, f)
        for f in os.listdir(PDF_DIR)
        if f.lower().endswith((".pdf", ".docx"))
    ]
    
    if not files:
        print(f"No PDF or DOCX files found in {PDF_DIR}")
        return

    print(f"Found {len(files)} files to process...")
    
    texts = []
    metadatas = []
    
    for file_idx, file in enumerate(files):
        print(f"Processing file {file_idx + 1}/{len(files)}: {os.path.basename(file)}")
        
        if file.lower().endswith(".pdf"):
            raw_text = extract_pdf_text(file)
        elif file.lower().endswith(".docx"):
            raw_text = extract_docx_text(file)
        else:
            continue

        if not raw_text.strip():
            print(f"  Warning: No text extracted from {file}")
            continue
            
        # Split into overlapping chunks for effective vector search
        chunk_size = 700
        overlap = 100
        chunks = []
        
        for i in range(0, len(raw_text), chunk_size - overlap):
            chunk = raw_text[i:i+chunk_size]
            if chunk.strip():
                chunks.append(chunk.strip())
                metadatas.append({
                    "source_file": os.path.basename(file),
                    "chunk_index": len(chunks) - 1
                })
        
        texts.extend(chunks)
        print(f"  Extracted {len(chunks)} chunks from {os.path.basename(file)}")

    if texts:
        print(f"\nGenerating embeddings for {len(texts)} text chunks...")
        embeddings = embed_model.encode(texts).tolist()
        
        print("Adding to database...")
        db.add(
            documents=texts, 
            embeddings=embeddings, 
            metadatas=metadatas,
            ids=[str(i) for i in range(len(texts))]
        )
        
        # Verify ingestion
        final_count = db.count()
        print(f"✅ Successfully ingested {final_count} chunks from {len(files)} files into 'adgm_laws'.")
        print(f"Database stored at: {DB_PATH}")
        
        # Test a sample query
        test_results = db.query(query_texts=["ADGM company"], n_results=1)
        if test_results['documents'][0]:
            print(f"✅ Database test query successful: {test_results['documents'][0][0][:50]}...")
        else:
            print("⚠️  Database test query returned no results")
            
    else:
        print("❌ No text found to ingest.")

if __name__ == "__main__":
    ingest()